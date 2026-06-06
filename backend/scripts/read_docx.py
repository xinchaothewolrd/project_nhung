import os
import docx

doc_path = r"c:\Users\Lenovo\Desktop\do an\mon t4\yeu cau\Phân công công việc.docx"
if os.path.exists(doc_path):
    doc = docx.Document(doc_path)
    fullText = []
    for para in doc.paragraphs:
        fullText.append(para.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                fullText.append(cell.text)
    text = "\n".join(fullText)
    
    with open("extracted_requirements.txt", "w", encoding="utf-8") as f:
        f.write(text)
    print("Success")
else:
    print("Docx file not found")
